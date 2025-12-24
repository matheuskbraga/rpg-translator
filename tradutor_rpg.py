import customtkinter as ctk
from tkinter import filedialog, messagebox
import fitz  # PyMuPDF
from deep_translator import GoogleTranslator
from docx import Document
import threading
import os
import time

# Configuração inicial do tema
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class FalloutTranslatorApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("Tradutor de PDF RPG (Fallout Edition)")
        self.geometry("600x450")
        self.resizable(False, False)

        # Variáveis de controle
        self.file_path = None
        self.is_running = False

        # Layout
        self.create_widgets()

    def create_widgets(self):
        # Título
        self.label_title = ctk.CTkLabel(self, text="Tradutor de PDF para PT-BR", font=("Roboto", 24, "bold"))
        self.label_title.pack(pady=20)

        # Botão de Seleção de Arquivo
        self.btn_select = ctk.CTkButton(self, text="Selecionar Arquivo PDF", command=self.select_file)
        self.btn_select.pack(pady=10)

        # Label do arquivo selecionado
        self.label_file = ctk.CTkLabel(self, text="Nenhum arquivo selecionado", text_color="gray")
        self.label_file.pack(pady=5)

        # Barra de Progresso
        self.progress_bar = ctk.CTkProgressBar(self, width=400)
        self.progress_bar.set(0)
        self.progress_bar.pack(pady=20)

        # Label de Status
        self.label_status = ctk.CTkLabel(self, text="Aguardando...", font=("Roboto", 14))
        self.label_status.pack(pady=5)

        # Botão de Iniciar Tradução
        self.btn_start = ctk.CTkButton(self, text="Iniciar Tradução", command=self.start_thread, state="disabled", fg_color="green")
        self.btn_start.pack(pady=20)

        # Console de Log (Textbox)
        self.textbox_log = ctk.CTkTextbox(self, width=500, height=120)
        self.textbox_log.pack(pady=10)
        self.textbox_log.insert("0.0", "Logs do sistema aparecerão aqui...\n")

    def select_file(self):
        filetypes = (("PDF files", "*.pdf"), ("All files", "*.*"))
        filename = filedialog.askopenfilename(title="Selecione o PDF do Fallout", initialdir="/", filetypes=filetypes)
        
        if filename:
            self.file_path = filename
            self.label_file.configure(text=f"Arquivo: {os.path.basename(filename)}")
            self.btn_start.configure(state="normal")
            self.log(f"Arquivo carregado: {filename}")

    def log(self, message):
        self.textbox_log.insert("end", message + "\n")
        self.textbox_log.see("end")

    def start_thread(self):
        if not self.file_path:
            return
        
        self.is_running = True
        self.btn_select.configure(state="disabled")
        self.btn_start.configure(state="disabled")
        
        # Inicia o processo em uma thread separada para não travar a GUI
        thread = threading.Thread(target=self.process_pdf)
        thread.start()

    def split_text(self, text, max_chars=4500):
        """Divide o texto em pedaços menores para respeitar limites da API"""
        return [text[i:i+max_chars] for i in range(0, len(text), max_chars)]

    def process_pdf(self):
        try:
            doc = fitz.open(self.file_path)
            total_pages = len(doc)
            
            # Cria documento Word para saída
            output_doc = Document()
            output_doc.add_heading(f'Tradução: {os.path.basename(self.file_path)}', 0)

            translator = GoogleTranslator(source='auto', target='pt')
            
            self.log(f"Iniciando tradução de {total_pages} páginas...")

            for i, page in enumerate(doc):
                if not self.is_running:
                    break

                page_num = i + 1
                self.label_status.configure(text=f"Traduzindo página {page_num} de {total_pages}...")
                self.log(f"Processando pág. {page_num}...")

                # Extrai texto da página
                text = page.get_text()
                
                if text.strip():
                    output_doc.add_heading(f'Página {page_num}', level=1)
                    
                    # Limpeza básica (remover quebras de linha excessivas em meio a frases)
                    text = text.replace('\n', ' ')
                    
                    # Traduzir em chunks (pedaços) para não estourar limite de caracteres
                    chunks = self.split_text(text)
                    translated_full = ""
                    
                    for chunk in chunks:
                        try:
                            translated_chunk = translator.translate(chunk)
                            translated_full += translated_chunk + " "
                            time.sleep(0.5) # Delay para evitar bloqueio de IP (Rate Limit)
                        except Exception as e:
                            self.log(f"Erro ao traduzir trecho na pág {page_num}: {e}")
                            translated_full += "[Erro na tradução deste trecho] "

                    output_doc.add_paragraph(translated_full)
                
                # Atualiza barra de progresso
                progress = page_num / total_pages
                self.progress_bar.set(progress)

            # Salvar arquivo final
            save_path = self.file_path.replace(".pdf", "_PT-BR.docx")
            output_doc.save(save_path)
            
            self.label_status.configure(text="Concluído!")
            self.log(f"Tradução salva em: {save_path}")
            messagebox.showinfo("Sucesso", f"Tradução concluída!\nSalvo em: {save_path}")

        except Exception as e:
            self.log(f"Erro fatal: {str(e)}")
            messagebox.showerror("Erro", str(e))
        
        finally:
            self.is_running = False
            self.btn_select.configure(state="normal")
            self.btn_start.configure(state="normal")

if __name__ == "__main__":
    app = FalloutTranslatorApp()
    app.mainloop()